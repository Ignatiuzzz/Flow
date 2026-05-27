from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


WORD_EXPORT_DIR = Path("app/uploads/exports/word")


def ensure_word_directory() -> None:
    WORD_EXPORT_DIR.mkdir(parents=True, exist_ok=True)


def safe_text(value: Optional[object]) -> str:
    if value is None:
        return ""

    return str(value)


def build_word_filename(prefix: str, code: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    clean_code = code.replace("/", "-").replace("\\", "-").replace(" ", "_")

    return f"{prefix}_{clean_code}_{timestamp}.docx"


def add_title(document: Document, title: str) -> None:
    paragraph = document.add_heading(title, level=1)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_field_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Campo"
    header_cells[1].text = "Detalle"

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = safe_text(row[0])
        cells[1].text = safe_text(row[1])

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def create_basic_word(
    filename: str,
    title: str,
    rows: list[list[str]]
) -> str:
    ensure_word_directory()

    file_path = WORD_EXPORT_DIR / filename

    document = Document()

    add_title(document, title)
    document.add_paragraph("")

    add_field_table(document, rows)

    document.save(file_path)

    return str(file_path).replace("\\", "/")


def generate_finding_word(finding: dict) -> str:
    filename = build_word_filename(
        prefix="ficha_hallazgo",
        code=safe_text(finding.get("codigo", "sin_codigo"))
    )

    rows = [
        ["Código", finding.get("codigo")],
        ["Nombre", finding.get("nombre")],
        ["Descripción", finding.get("descripcion")],
        ["Criterio", finding.get("criterio")],
        ["Objetivo", finding.get("objetivo")],
        ["Causa", finding.get("causa")],
        ["Efecto", finding.get("efecto")],
        ["Conclusión", finding.get("conclusion")],
        ["Impacto", finding.get("impacto")],
        ["Urgencia", finding.get("urgencia")],
        ["Riesgo", finding.get("riesgo")],
        ["Nivel", finding.get("nivel")],
        ["Justificación del Riesgo", finding.get("justificacionRiesgo")],
        ["Recomendaciones", finding.get("recomendaciones")],
    ]

    return create_basic_word(
        filename=filename,
        title="Ficha de Hallazgo",
        rows=rows
    )


def generate_evidence_word(evidence: dict) -> str:
    filename = build_word_filename(
        prefix="ficha_evidencia",
        code=safe_text(evidence.get("codigo", "sin_codigo"))
    )

    rows = [
        ["Código", evidence.get("codigo")],
        ["Nombre", evidence.get("nombre")],
        ["Criterio", evidence.get("criterio")],
        ["Objetivo", evidence.get("objetivo")],
        ["Descripción de Evidencia", evidence.get("descripcionEvidencia")],
        ["Documento", evidence.get("documentoNombre")],
        ["Subtítulo", evidence.get("subtitulo")],
        ["Hallazgo Relacionado", evidence.get("hallazgoId")],
    ]

    return create_basic_word(
        filename=filename,
        title="Ficha de Evidencia",
        rows=rows
    )